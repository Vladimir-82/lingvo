"""Сравнение оригинального текста и перевода."""

from docx import Document

from exceptions import NameErrorException
from lang.models import Translate
from lang.structures import Language


def get_document(translate_id: int) -> tuple[Document, str]:
    """Создание отчета."""
    translate_data = get_translate(translate_id)
    document = Document()
    document.add_heading(f"Перевод текста {translate_data.get('date')}", level=2)

    table = document.add_table(rows=3, cols=2)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Текст оригинала'
    hdr_cells[1].text = 'Перевод'

    lang = table.rows[1].cells
    lang[0].text = Language.language.get(translate_data.get('language_input'))
    lang[1].text = Language.language.get(translate_data.get('language_output'))

    text = table.rows[2].cells
    text[0].text = translate_data.get('translate_from_text')
    text[1].text = translate_data.get('translated_text')

    return document, get_file_name(translate_data)


def get_translate(translate_id: int) -> dict[str, str]:
    """Получение перевода для сравнения."""
    translate = get_translate_instance(translate_id)
    translate_data = {
        'language_input': translate.language_input,
        'language_output': translate.language_output,
        'translate_from_text': translate.translate_from_text,
        'translated_text': translate.translated_text,
        'date': translate.created_at.strftime('%d.%m.%Y %H:%M'),
    }
    return translate_data


def get_translate_instance(translate_id: int) -> Translate:
    """Получение объекта перевода."""
    return Translate.objects.get(id=translate_id)


def get_file_name(translate_data: dict[str, str]) -> str:
    """Получение имени файла отчета."""
    try:
        name = translate_data.get('translate_from_text')[:15]
    except TypeError:
        raise NameErrorException('Невозможно создать имя для переода.')
    return name
